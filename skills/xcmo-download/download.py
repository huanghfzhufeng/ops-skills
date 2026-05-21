#!/usr/bin/env python3
"""xcmo-download: 从 xcmo 平台批量下载 batch 产物。

用法:
    python3 download.py --external "batch-A,batch-B" --internal "batch-X"
    python3 download.py --internal "batch-only-internal"
    python3 download.py --external "batch-only-external"
    python3 download.py --external "batch-A" --out-dir /tmp/custom-out

读取 ~/.claude/memory/xcmo-session.json 拿 vee_session token 做认证。

输出位置：~/Desktop/xcmo-batches/{YYYYMMDD}/
命名规则：
    {YYYYMMDD} 南宁合作方.docx  + .zip   （外部）
    {YYYYMMDD} 内部.docx                  （内部，不打 zip）

退出码：
    0  成功
    2  python-docx 缺失
    3  session 文件缺失
    4  TOKEN 失效（401/403）—— Claude 应引导用户更新 token
"""

import argparse
import json
import sys
import zipfile
from datetime import datetime
from pathlib import Path
import urllib.request
import urllib.error

try:
    from docx import Document
except ImportError:
    print("❌ 缺少 python-docx, 请先安装: pip install python-docx", file=sys.stderr)
    sys.exit(2)

XCMO_BASE = "https://xcmo.ai"
SESSION_FILE = Path.home() / ".claude/memory/xcmo-session.json"

# Cloudflare 拒绝默认 Python-urllib UA，必须用浏览器 UA 绕开 Error 1010
USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)

# 类别命名（外部 vs 内部）
EXTERNAL_LABEL = "南宁合作方"
INTERNAL_LABEL = "内部"


class AuthExpiredError(RuntimeError):
    """xcmo session token 失效或过期，需要用户提供新 token。"""


def load_session() -> dict:
    if not SESSION_FILE.exists():
        print(
            f"❌ 没找到 xcmo session 文件: {SESSION_FILE}\n"
            "请先告诉 Claude 你的 vee_session token，让它写入这个文件。",
            file=sys.stderr,
        )
        sys.exit(3)
    data = json.loads(SESSION_FILE.read_text(encoding="utf-8"))
    if not data.get("vee_session"):
        print(
            f"❌ session 文件存在但 vee_session 字段为空: {SESSION_FILE}\n"
            "请告诉 Claude 更新 token：'更新 xcmo token: <你的 token>'",
            file=sys.stderr,
        )
        sys.exit(3)
    return data


def http_get(url: str, session: dict, expect_json: bool = True):
    """GET with cookie auth. 401/403 抛 AuthExpiredError 让上层引导用户。"""
    req = urllib.request.Request(
        url,
        headers={
            "Cookie": f"vee_session={session['vee_session']}",
            "Accept": "application/json",
            "User-Agent": USER_AGENT,
        },
    )
    try:
        resp = urllib.request.urlopen(req, timeout=30)
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")[:300]
        if e.code in (401, 403):
            raise AuthExpiredError(
                f"HTTP {e.code}: token 失效或权限不足。body: {body}"
            )
        raise RuntimeError(f"HTTP {e.code} {url}\n  body: {body}")
    if expect_json:
        return json.loads(resp.read().decode("utf-8"))
    return resp.read()


def download_file(url: str, dest: Path, session: dict) -> bool:
    """下载文件到本地，返回是否成功。失败重试 1 次。"""
    full_url = url if url.startswith("http") else XCMO_BASE + url
    req = urllib.request.Request(
        full_url,
        headers={
            "Cookie": f"vee_session={session['vee_session']}",
            "User-Agent": USER_AGENT,
        },
    )
    for attempt in range(2):
        try:
            with urllib.request.urlopen(req, timeout=180) as resp:
                dest.write_bytes(resp.read())
            return True
        except urllib.error.HTTPError as e:
            if e.code in (401, 403):
                raise AuthExpiredError(f"HTTP {e.code} 下载视频时 token 失效")
            if attempt == 0:
                print(f"    ⚠ 第 1 次下载失败 ({e}), 重试…", file=sys.stderr)
            else:
                print(f"    ❌ 下载失败 {full_url}: {e}", file=sys.stderr)
        except Exception as e:
            if attempt == 0:
                print(f"    ⚠ 第 1 次下载失败 ({e}), 重试…", file=sys.stderr)
            else:
                print(f"    ❌ 下载失败 {full_url}: {e}", file=sys.stderr)
    return False


def fetch_batch_assets(batch_id: str, session: dict) -> dict:
    """获取一个 batch 的所有 task + asset 完整数据。"""
    print(f"  ◾ 查询 batch {batch_id}")
    batch_data = http_get(f"{XCMO_BASE}/api/tasks/batch/{batch_id}", session)
    items = []
    for task in batch_data.get("tasks", []):
        if task["status"] != "completed":
            print(f"    ⚠ task {task['id'][:8]} status={task['status']}, skip")
            continue
        asset_id = (task.get("result") or {}).get("asset_id")
        if not asset_id:
            print(f"    ⚠ task {task['id'][:8]} no asset_id, skip")
            continue
        try:
            asset = http_get(f"{XCMO_BASE}/api/assets/{asset_id}", session)
            items.append({"task": task, "asset": asset})
        except AuthExpiredError:
            raise
        except Exception as e:
            print(f"    ❌ asset {asset_id[:8]} fetch fail: {e}", file=sys.stderr)
    return {
        "batch_id": batch_id,
        "total": batch_data.get("total", 0),
        "completed": batch_data.get("completed", 0),
        "items": items,
    }


def sanitize_filename(name: str, max_len: int = 50) -> str:
    """文件名安全化（去掉特殊字符）。"""
    bad = '<>:"/\\|?*\n\r\t'
    cleaned = "".join(c if c not in bad else "-" for c in (name or "")).strip()
    return cleaned[:max_len] if len(cleaned) > max_len else cleaned


def video_filename(asset: dict) -> str:
    """生成本地视频文件名：{YYYYMMDD} {character} {任务名} {asset8}.mp4

    日期来自 asset.created_at（视频生成时间，非下载时间），方便排序和追溯。
    asset 短 ID 防同名冲突，也能反查 xcmo 系统里的具体 asset。
    """
    name = sanitize_filename(asset.get("name") or "video")
    char = sanitize_filename(asset.get("character_id") or "unknown")
    created = asset.get("created_at") or ""
    # 解析 ISO datetime: "2026-05-20T11:13:01.452965" → "20260520"
    date_str = created[:10].replace("-", "") if len(created) >= 10 else "00000000"
    short_id = asset["id"][:8]
    return f"{date_str} {char} {name} {short_id}.mp4"


def build_docx(category_label: str, batches_data: list, out_path: Path, video_files: dict) -> int:
    """生成 Word 文档，返回成功导出的视频数量。"""
    doc = Document()

    doc.add_heading(f"xcmo 批次产物 — {category_label}", level=0)

    total_videos = sum(len(b["items"]) for b in batches_data)
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
    meta_p.add_run(f"批次数: {len(batches_data)}  |  视频数: {total_videos}").italic = True

    for batch_info in batches_data:
        doc.add_heading(f"batch: {batch_info['batch_id']}", level=1)
        sub = doc.add_paragraph()
        sub.add_run(
            f"task 总数: {batch_info.get('total', 0)}  |  "
            f"成功导出: {len(batch_info['items'])}"
        ).italic = True

        if batch_info.get("error"):
            err_p = doc.add_paragraph()
            err_p.add_run(f"⚠ 拉取失败: {batch_info['error']}").bold = True
            continue

        if not batch_info["items"]:
            doc.add_paragraph("（无可导出的 task）")
            continue

        for i, item in enumerate(batch_info["items"], 1):
            asset = item["asset"]

            video_name = asset.get("name") or f"视频 {i}"
            doc.add_heading(f"{i}. {video_name}", level=2)

            # 视频文件名
            fn_p = doc.add_paragraph()
            fn_p.add_run("视频文件: ").bold = True
            local_video_name = video_filename(asset)
            fn_p.add_run(local_video_name)
            if asset["id"] in video_files and not video_files[asset["id"]]:
                fn_p.add_run("  ⚠ 视频下载失败").bold = True

            # 角色
            ch_p = doc.add_paragraph()
            ch_p.add_run("角色: ").bold = True
            ch_p.add_run(asset.get("character_id") or "—")

            # Caption
            cap_p = doc.add_paragraph()
            cap_p.add_run("Caption: ").bold = True
            cap_p.add_run(asset.get("caption") or "—")

            # Hashtags
            tag_p = doc.add_paragraph()
            tag_p.add_run("Hashtags: ").bold = True
            tags = asset.get("asset_hashtags") or []
            tag_p.add_run(" ".join(tags) if tags else "—")

    doc.save(out_path)
    return total_videos


def build_zip(batches_data: list, video_files: dict, out_path: Path) -> int:
    """打包视频成 zip，返回打包的视频数量。"""
    count = 0
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for batch_info in batches_data:
            for item in batch_info["items"]:
                asset = item["asset"]
                local_path = video_files.get(asset["id"])
                if local_path and isinstance(local_path, Path) and local_path.exists():
                    # 平铺打包，不要 batch 子目录（文件名已含日期+角色+任务名足够识别）
                    arc_name = local_path.name
                    zf.write(local_path, arc_name)
                    count += 1
    return count


def process_category(
    category_label: str,
    batch_ids: list,
    work_dir: Path,
    session: dict,
    date_str: str,
    *,
    pack_zip: bool,
):
    """处理一个类别（外部/内部）的所有 batch。

    pack_zip=True  → 同时生成 .docx + .zip （外部用，方便发合作方）
    pack_zip=False → 只生成 .docx，视频留 videos/ 子目录（内部用，自己拿）
    """
    if not batch_ids:
        print(f"⏭  {category_label}: 没有 batch，跳过")
        return None

    print(f"\n📦 处理 {category_label}（{len(batch_ids)} 个 batch）")
    batches_data = []
    video_files = {}  # asset_id -> Path (or None if 下载失败)

    for bid in batch_ids:
        bid = bid.strip()
        if not bid:
            continue
        try:
            batch_info = fetch_batch_assets(bid, session)
        except AuthExpiredError:
            raise
        except Exception as e:
            print(f"  ❌ batch {bid} 拉取失败: {e}", file=sys.stderr)
            batches_data.append({"batch_id": bid, "total": 0, "completed": 0, "items": [], "error": str(e)})
            continue

        for item in batch_info["items"]:
            asset = item["asset"]
            file_url = asset.get("file_url")
            if not file_url:
                print(f"    ⚠ asset {asset['id'][:8]} 无 file_url, skip 下载")
                video_files[asset["id"]] = None
                continue
            local_filename = video_filename(asset)
            local_path = work_dir / "videos" / category_label / local_filename
            local_path.parent.mkdir(parents=True, exist_ok=True)
            print(f"    ⬇ 下载 {asset.get('name', '')[:30]}...")
            ok = download_file(file_url, local_path, session)
            video_files[asset["id"]] = local_path if ok else None

        batches_data.append(batch_info)

    # 生成 docx
    docx_path = work_dir / f"{date_str} {category_label}.docx"
    n_videos = build_docx(category_label, batches_data, docx_path, video_files)
    print(f"  ✓ 文档: {docx_path.name}")

    result = {
        "category": category_label,
        "batches": len(batches_data),
        "videos_total": n_videos,
        "docx": str(docx_path),
        "failed_batches": [b["batch_id"] for b in batches_data if b.get("error")],
    }

    # 只有外部生成 zip
    if pack_zip:
        zip_path = work_dir / f"{date_str} {category_label}.zip"
        n_zipped = build_zip(batches_data, video_files, zip_path)
        print(f"  ✓ 视频包: {zip_path.name}（{n_zipped} 个视频）")
        result["zip"] = str(zip_path)
        result["videos_zipped"] = n_zipped
    else:
        videos_dir = work_dir / "videos" / category_label
        n_local = sum(1 for v in video_files.values() if v and isinstance(v, Path) and v.exists())
        print(f"  ✓ 视频文件: {videos_dir}/ （{n_local} 个，不打包）")
        result["videos_dir"] = str(videos_dir)
        result["videos_saved"] = n_local

    return result


def parse_csv_list(s: str) -> list:
    if not s:
        return []
    return [x.strip() for x in s.split(",") if x.strip()]


def main():
    parser = argparse.ArgumentParser(description="xcmo batch downloader")
    parser.add_argument("--external", default="", help="外部（南宁合作方）batch IDs, 逗号分隔")
    parser.add_argument("--internal", default="", help="内部 batch IDs, 逗号分隔")
    parser.add_argument("--out-dir", default="", help="输出目录（默认 ~/Desktop/xcmo-batches/<YYYYMMDD>/）")
    args = parser.parse_args()

    external = parse_csv_list(args.external)
    internal = parse_csv_list(args.internal)

    if not external and not internal:
        sys.exit("❌ 没有提供 batch ID。用 --external 或 --internal")

    now = datetime.now()
    date_str = now.strftime("%Y%m%d")
    work_dir = Path(args.out_dir) if args.out_dir else Path.home() / "Desktop" / "xcmo-batches" / date_str
    work_dir.mkdir(parents=True, exist_ok=True)

    session = load_session()
    print(f"📁 输出目录: {work_dir}")
    print(f"🔑 认证: {session.get('user_email')}")
    print(f"🆔 user_id: {session.get('user_id')}")
    print(f"📅 日期: {date_str}")

    summary = {
        "out_dir": str(work_dir),
        "date_str": date_str,
        "external_count": len(external),
        "internal_count": len(internal),
        "results": [],
    }

    try:
        # 外部 → docx + zip
        result_ext = process_category(
            EXTERNAL_LABEL, external, work_dir, session, date_str, pack_zip=True
        )
        if result_ext:
            summary["results"].append(result_ext)
        # 内部 → 只 docx
        result_int = process_category(
            INTERNAL_LABEL, internal, work_dir, session, date_str, pack_zip=False
        )
        if result_int:
            summary["results"].append(result_int)
    except AuthExpiredError as e:
        print("\n" + "=" * 60, file=sys.stderr)
        print("❌ xcmo session token 失效或过期", file=sys.stderr)
        print(f"   错误: {e}", file=sys.stderr)
        print("", file=sys.stderr)
        print("解决：让 Claude 帮你更新 token：", file=sys.stderr)
        print("  1. 浏览器打开 https://xcmo.ai 登录", file=sys.stderr)
        print("  2. F12 → Application → Cookies → 复制 vee_session 值", file=sys.stderr)
        print("  3. 告诉 Claude：'更新 xcmo token: <你复制的 token>'", file=sys.stderr)
        print("  4. Claude 会写入 ~/.claude/memory/xcmo-session.json，然后重试下载", file=sys.stderr)
        sys.exit(4)

    print("\n" + "=" * 60)
    print("✅ 完成")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
